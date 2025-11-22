# doc_auth_checker.py  (ABSTRACT STRUCTURE)

# --- Imports ---
import os
import datetime
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from difflib import SequenceMatcher
from langdetect import detect
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# --- Helper functions (metadata + text) ---
def similar(a, b):
    """Return similarity ratio between two strings."""
    return SequenceMatcher(None, a, b).ratio()

def get_filesystem_metadata(file_path):
    """Extract filesystem created/modified timestamps."""
    stat = os.stat(file_path)
    try:
        created = datetime.datetime.fromtimestamp(stat.st_birthtime)
    except AttributeError:
        created = datetime.datetime.fromtimestamp(stat.st_ctime)
    modified = datetime.datetime.fromtimestamp(stat.st_mtime)
    return created, modified

def get_docx_metadata(file_path):
    """Extract Word metadata: author, created, modified, language."""
    try:
        doc = Document(file_path)
        prop = doc.core_properties
        metadata = {
            "author": prop.author,
            "created": prop.created,
            "modified": prop.modified,
            "language": prop.language,
        }
        return metadata
    except PackageNotFoundError:
        return None

def read_file_content(file_path):
    """Read document paragraphs and return as single string."""
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except PackageNotFoundError:
        return None

# --- Similarity engine ---
def compare_documents(directory, main_file_path):
    """Compare content similarity with all other documents in directory."""
    from doc_auth_checker import read_file_content  # local import
    documents = []
    file_names = []

    main_document_content = read_file_content(main_file_path)
    documents.append(main_document_content)

    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)

        if file_path == main_file_path or not file_path.endswith('.docx'):
            continue

        content = read_file_content(file_path)
        if content:
            documents.append(content)
            file_names.append(file_name)

    tfidf_vectorizer = TfidfVectorizer()
    tfidf_matrix = tfidf_vectorizer.fit_transform(documents)
    cosine_similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

    similar_documents = [
        (file_names[i], cosine_similarities[i])
        for i in range(len(file_names))
        if cosine_similarities[i] > 0.58
    ]
    return similar_documents

def analyze_document(file_path, student_name, published_date, due_date, directory=r"D:\documents"):
    """Perform authenticity checks: metadata + timestamps + similarity."""
    from doc_auth_checker import get_filesystem_metadata, get_docx_metadata, similar
    suspect_events = []

    fs_created, fs_modified = get_filesystem_metadata(file_path)
    metadata = get_docx_metadata(file_path)

    if not metadata:
        return ["Unable to extract metadata: Unsupported or corrupted file"]

    if fs_created < published_date:
        suspect_events.append("Filesystem creation date is before assignment publish date.")
    if fs_modified < published_date:
        suspect_events.append("Filesystem modification date is before assignment publish date.")
    if fs_modified > due_date:
        suspect_events.append("Filesystem modification date is after assignment due date.")

    author_similarity = similar(metadata['author'].lower(), student_name.lower())
    if author_similarity < 0.5:
        suspect_events.append(f"Low author similarity ratio: {author_similarity:.2f}")

    editing_time = metadata['modified'] - metadata['created']
    if editing_time > datetime.timedelta(hours=100):
        suspect_events.append(f"Editing time unusually long ({editing_time}).")
    if editing_time < datetime.timedelta(minutes=30):
        suspect_events.append(f"Editing time unusually short ({editing_time}).")

    similar_docs = compare_documents(directory, file_path)
    if similar_docs:
        for doc_name, sim_score in similar_docs:
            suspect_events.append(f"Similar to {doc_name} with score {sim_score:.2f}")

    return suspect_events